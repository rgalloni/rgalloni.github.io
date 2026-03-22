#!/usr/bin/env python3
"""
Combina i file Markdown dell'itinerario e genera un DOCX via CloudConvert.

Supporta placeholder immagine direttamente nei file .md, ad esempio:
[[IMAGE:images/01_hippodrome_of_constantinople.jpg]]
[[IMAGE:images/01_hippodrome_of_constantinople.jpg|width=2.2|align=left]]

I file Markdown sono la source of truth del contenuto.
"""

from __future__ import annotations

import argparse
import os
import re
import sys
import time
from pathlib import Path
from typing import Iterable, List

import requests
from dotenv import load_dotenv


load_dotenv()

DEFAULT_ITINERARY_DIR = Path("itinerary")
DEFAULT_COMBINED_MD_NAME = "_combined_for_cloudconvert.md"
DEFAULT_OUTPUT_DOCX = "Itinerario_Turchia_Istanbul_Giorni_1_2_CloudConvert.docx"
API_KEY = os.environ.get("CLOUDCONVERT_API_KEY")

IMAGE_PLACEHOLDER_RE = re.compile(
    r"^\[\[IMAGE:(?P<path>[^\]|]+?)(?:\|width=(?P<width>\d+(?:\.\d+)?))?(?:\|align=(?P<align>left|center|right))?\]\]$"
)


def check_dependencies() -> None:
    try:
        import docx  # noqa: F401
        import requests  # noqa: F401
        from dotenv import load_dotenv as _load_dotenv  # noqa: F401
    except ImportError:
        print("❌ Mancano dipendenze richieste: requests, python-doten
I file Markdown sono la source of truth del contenuto.
"""

from __futuon-dotenv python-docx")
        sys.exit(1)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Converte i Markdown dell'itinerario in DOCX.")
    parser.add_argument(
        "--itinerary-dir",
        type=Path,
        default=DEFAULT_ITINERARY_DIR,
        help="Cartella contenenDEFAULT_OUTPUT_DOCX = "Itinerario_Turchia_Is    parser.add_argument(
        "--combined-name",
        default=DEFAULT_COMBINED_MD_NAME,
        help="Nome del file Markdown combinato.",
    )
    parser.add_argument(
      )


def check_dependencies() -> None:
    try:
        import docx  # noqa: F401
        import re)
    parser.add_argument(
        "--skip-cloudconvert",
        action="store_true",
        help="Genera solo il Markdown combinato, senza conversione CloudConvert.",
    )
    parser.add_argument(
        "--postprocess-docx",
        type=Path,
        help="Applica solo la sostituzione dei placehold"""

from __futuon-dotenv python-docx")
        sys.earser.parse_args()


def list_markdown_files(itinerary_dir: Path, combined_name: str) -> List[Path]:
       parser.add_argument(
        "--itinerary-dir",path in itinerary_dir.glob("*.md")
        if path.name != combined_name and not path.name.startswith("_")
    )
    if not markdown_files:
        raise FileNotFoundError(f"Nessun file Markdown trovato in {itinerary_dir}")
    return markdown_files


def combine_markdown_files(markdown_files: Iterable[Path], combined_path: Path) -> Path:
    files = list(markdown_files)
    page_break_code = '```{=openxml}\n<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```'

    with combined_path.open("w", encoding="utf-8") as outfile:
        for index, file_path in enumerate(files):
            outfile.write(file_path.read_text(encoding="utf-8").strip())
            if index < len(files) - 1:
                outfile.write(f"\n\n{page_break_code}\n\n")

    print(f"✅ File combinato creato: {combined_path}")
    return combined_path


def wait_for_job(job_id: str) -> dict:
    start_time = t       parser.add_argument(
  = 180

    while True:
        response = requests.get(
            f"https://api.cloudconvert.com/v2/jobs/{job_id}",
            headers={"Authorization": f"Bearer {API_KEY}"},
            timeout=60,
        ).json()

        status = response["data"]["status"]
        if status == "finished":
            return response
        if status == "error":
     files = list(markdown_files)
    page_break_code = '```{=openxml}\n<w:p><w:r><w:br int(f"❌ Errore CloudConvert: {message}")
            sys.exit(1)
        if time.time() - start_time > max_wait_seconds:
            print(f"❌ Timeout CloudConvert dopo {max_wait_seconds} secondi")
            sys.exit(1)

        pr           o job: {status}...")
        time.sleep(3)


def convert_with_cloudconvert(markdown_path: Path, output_path: Path) -> None:
    if not API_KEY or API_KEY == "LA_TUA_CHIAVE_API_QUI":
        print("❌ CLOUDCONVERT_API_KEY non impostata nel file .env")
        sys.exit(1)

    payload = {
        "tasks": {
            "import-file": {"operation": "import/upload"},
            "convert-file": {
                "operation": "convert",
                "input": "import-file",
                "input_format": "md",
                "output_format": "docx",
                "engine": "pandoc",
            },
            "export-file": {
                "operation": "export/url",
                "input": "convert-file",
                "inline": False,
                "archive_multiple_files": False,
            },
          if time.time()nse = requests.post(
        "https://api.cloudconvert.com/v2/jobs",
        headers={
            "Authorization": f"Bearer {API_KEY}",
            "Content-type": "application/json",
        },
        json=payload,
        timeout=60,
    ).json()

    if "data" not in response:
        print(f"❌ Errore creazione job CloudConvert: {response}")
        sys.exit(1)

    job_id = response["data"]["id"]
    upload_url = response["data"]["tasks"][0]["result"]["form"]["url"]
    upload_params = response["data"]["tasks"][0]["result"]["form"]["parameters"]

    print(f"   - Job creato: {job_id}")
    with markdown_path.open("rb") as file_handle:
        upload_response = requests.post(
            upload_url,
            data=upload_params,
            files={"file": file_handle},
                 "operation"   )

    if upload_response.status_code != 201:
        print(f"❌ Errore upload file: {upload_response.text}")
        sys.exit(1)

    job_data = wait_for_job(job_id)
    export_task = next(task for task in job_data["data"]["tasks"] if task["operation"] == "export/url")
    download_url = export_task["result"]["files"][0]["url"]

    downloaded_file = requests.get(download_url, timeout=120)
    downloaded_file.raise_for_status()
    output_path.write_bytes(downloaded_file.content)
    print(f"✅ DOCX scaricato: {output_path}")


def parse_image_placeholder(text: str):
    stripped = text.strip()
    if not (stripped.startswith("[[IMAGE:") and stripped.endswith("]]")):
        return None

    inner = stripped[len("[[IMAGE:"):-2]
    parts = [part.strip() for part in inner.split("|") if part.strip()]
    if not parts:
        return None

    image_path = parts[0]
    width = 2.2
              data  for part in parts[1:]:
        key, sep, value = part.partition("=")
        if not sep:
            continue
        key = key.strip().lower()
        value = value.strip()
        if key == "width":
            try:
                width = float(value)
            except ValueError:
                pass
        elif key == "align" and value in {"left", "center", "right"}:
            align = value

    return {"path": image_path, "width": width, "align": align}


def replace_image_placeholders_in_docx(docx_path: Path, base_dir: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    if not docx_path.exists():
        print(f"❌ DOCX non trovato: {docx_path}")
        sys.exit(1)

    document = Document(docx_path)
    replaced = 0

    for paragraph in list(document.paragraphs):
        spec = parse_image_placeholder(paragraph.text)
        if not
    image_path =continue

        image_path = (base_dir / spec["path"]).resolve()
        if not image_path.exists():
            print(f"⚠️ Immagine non trovata: {image_path}")
            continue

        paragraph.text = ""
        if spec["align"] == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif spec["align"] == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Inches(0.15)

        paragraph.add_run().add_picture(str(image_path), width=Inches(spec["width"]))
        replaced += 1

    document.save(docx_path)
    print(f"✅ Placeholder immagini sostituiti: {replaced}")


def main() -> None:
    check_dependencies()
    args = parse_args()

    itinerary_dir = args.itinerary_dir
    combined_path = itindir / args.combined_name
    output_path = Path(args.output_docx)

    if args.postprocess_docx:
        replace_image_placeholders_in_docx(args.postprocess_docx, itinerary_dir)
        return

    markdown_files = list_markdown_files(itinerary_dir, args.combined_name)
    combine_markdown_files(markdown_files, combined_path)

    if args.skip_cloudconvert:
        print("⏭️ CloudConvert saltato (--skip-cloudconvert)")
        return

    conve            paragraph.alignment = WD_t_path)
    replace_image_placeholders_in_docx(output_path, itinerary_dir)


if __name__ == "__main__":
    main()
