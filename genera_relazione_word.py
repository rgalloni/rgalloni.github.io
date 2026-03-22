#!/usr/bin/env python3
"""
Script per generare un documento Word completo della Relazione Viaggio 4808
Include tutte le sezioni della relazione e il manuale operativo del coordinatore
"""

import os
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def setup_document_styles(doc):
    """Configura gli stili del documento"""
    
    # Stile per il titolo principale
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(24)
    title_font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(24)
    
    # Stile per i titoli di sezione
    section_style = doc.styles.add_style('SectionTitle', WD_STYLE_TYPE.PARAGRAPH)
    section_font = section_style.font
    section_font.name = 'Arial'
    section_font.size = Pt(18)
    section_font.bold = True
    section_style.paragraph_format.space_before = Pt(18)
    section_style.paragraph_format.space_after = Pt(18) # Aumentato per maggiore separazione
    
    # Stile per i sottotitoli (es. Programma della Giornata)
    subtitle_style = doc.styles.add_style('SubTitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Arial'
    subtitle_font.size = Pt(14)
    subtitle_font.bold = True
    subtitle_style.paragraph_format.space_before = Pt(12)
    subtitle_style.paragraph_format.space_after = Pt(8)
    
    # Stile per i titoli di terzo livello (es. Mattina, Pomeriggio)
    heading3_style = doc.styles.add_style('Heading3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_font = heading3_style.font
    heading3_font.name = 'Arial'
    heading3_font.size = Pt(12)
    heading3_font.bold = True
    heading3_style.paragraph_format.space_before = Pt(10)
    heading3_style.paragraph_format.space_after = Pt(4)
    
    # Stile per il testo normale
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)

    # Definisci stili per elenchi puntati indentati
    try:
        # Livello 2
        list_bullet_2_style = doc.styles.add_style('List Bullet 2', WD_STYLE_TYPE.PARAGRAPH)
        list_bullet_2_style.base_style = doc.styles['List Bullet']
        list_bullet_2_style.paragraph_format.left_indent = Inches(0.5)
        
        # Livello 3
        list_bullet_3_style = doc.styles.add_style('List Bullet 3', WD_STYLE_TYPE.PARAGRAPH)
        list_bullet_3_style.base_style = doc.styles['List Bullet']
        list_bullet_3_style.paragraph_format.left_indent = Inches(1.0)
    except Exception:
        # Gli stili potrebbero già esistere se lo script viene eseguito più volte
        # in un contesto interattivo o se il template di base li ha.
        # In questo caso, non facciamo nulla.
        pass

def add_page_break(doc):
    """Aggiunge un'interruzione di pagina"""
    doc.add_page_break()

def process_markdown_content(doc, content, is_main_title=False):
    """Processa il contenuto markdown e lo aggiunge al documento Word"""
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Controlla se è una tabella (riga con |)
        if '|' in line and not line.startswith('#'):
            table_lines = []
            # Raccogli tutte le righe della tabella
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            
            # Processa la tabella
            if len(table_lines) >= 2:  # Almeno header e separatore
                process_table(doc, table_lines)
            continue
            
        # Titolo principale (# )
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip()
            if is_main_title:
                p = doc.add_paragraph(title_text, style='CustomTitle')
            else:
                p = doc.add_paragraph(title_text, style='SectionTitle')
        
        # Sottotitoli (## )
        elif line.startswith('## '):
            subtitle_text = line[3:].strip()
            doc.add_paragraph(subtitle_text, style='SubTitle')
        
        # Sottotitoli di terzo livello (### )
        elif line.startswith('### '):
            subtitle_text = line[4:].strip()
            doc.add_paragraph(subtitle_text, style='Heading3')
        
        # Liste con checkbox
        elif line.startswith('- [ ]'):
            list_text = line[5:].strip()
            p = doc.add_paragraph(f"☐ ", style='List Bullet')
            add_formatted_text(p, list_text)
        
        # Liste normali (con gestione indentazione)
        elif line.strip().startswith('- '):
            # Calcola il livello di indentazione basato sugli spazi
            indent_level = (len(line) - len(line.lstrip(' '))) // 2
            list_text = line.strip()[2:].strip()

            # Crea un paragrafo con lo stile base dell'elenco
            p = doc.add_paragraph(style='List Bullet')
            
            # Imposta il livello di indentazione corretto (ilvl)
            p_element = p._p
            pPr = p_element.get_or_add_pPr()
            numPr = pPr.get_or_add_numPr()
            
            # Imposta il livello di indentazione
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), str(indent_level))
            numPr.append(ilvl)
            
            # Imposta l'ID della numerazione (di solito 2 per i bullet di default)
            # Questo assicura che tutti i livelli usino lo stesso stile di punto elenco
            numId = OxmlElement('w:numId')
            numId.set(qn('w:val'), "2")
            numPr.append(numId)
            
            # Aggiungi il testo formattato al paragrafo
            add_formatted_text(p, list_text)
        
        # Testo normale
        else:
            # Rimuovi eventuali link markdown
            if '[' in line and '](' in line:
                import re
                line = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', line)
            
            p = doc.add_paragraph()
            add_formatted_text(p, line)
        
        i += 1

def add_formatted_text(paragraph, text, bold=False):
    """Aggiunge testo formattato al paragrafo gestendo grassetto, corsivo e altri stili"""
    import re
    
    if not text:
        return
    
    # Pattern per trovare formattazioni multiple
    # Ordine importante: prima grassetto+corsivo, poi singoli
    patterns = [
        (r'\*\*\*(.*?)\*\*\*', 'bold_italic'),  # ***testo*** = grassetto + corsivo
        (r'\*\*(.*?)\*\*', 'bold'),             # **testo** = grassetto
        (r'\*(.*?)\*', 'italic'),               # *testo* = corsivo
        (r'`(.*?)`', 'code')                    # `testo` = codice
    ]
    
    # Processa il testo con tutti i pattern
    current_text = text
    segments = [(current_text, 'normal')]
    
    for pattern, style in patterns:
        new_segments = []
        for segment_text, segment_style in segments:
            if segment_style == 'normal':
                parts = re.split(pattern, segment_text)
                for i, part in enumerate(parts):
                    if part:
                        if i % 2 == 0:
                            new_segments.append((part, 'normal'))
                        else:
                            new_segments.append((part, style))
            else:
                new_segments.append((segment_text, segment_style))
        segments = new_segments
    
    # Aggiungi i segmenti al paragrafo
    for segment_text, segment_style in segments:
        if not segment_text:
            continue
            
        run = paragraph.add_run(segment_text)
        
        # Applica gli stili
        if bold or segment_style in ['bold', 'bold_italic']:
            run.bold = True
        if segment_style in ['italic', 'bold_italic']:
            run.italic = True
        if segment_style == 'code':
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

def process_table(doc, table_lines):
    """Processa una tabella markdown e la converte in tabella Word"""
    if len(table_lines) < 2:
        return
    
    # Rimuovi la riga separatore (quella con ---)
    header_line = table_lines[0]
    data_lines = [line for line in table_lines[1:] if not re.match(r'^[\s\|\-:]+$', line)]
    
    if not data_lines:
        return
    
    # Conta le colonne dall'header
    header_cells = [cell.strip() for cell in header_line.split('|') if cell.strip()]
    num_cols = len(header_cells)
    
    if num_cols == 0:
        return
    
    # Crea la tabella
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    
    # Aggiungi header con formattazione
    header_row = table.rows[0]
    for i, cell_text in enumerate(header_cells):
        if i < len(header_row.cells):
            cell = header_row.cells[i]
            # Pulisci il paragrafo esistente
            cell.paragraphs[0].clear()
            # Aggiungi testo formattato
            add_formatted_text(cell.paragraphs[0], cell_text, bold=True)
    
    # Aggiungi righe dati con formattazione
    for line in data_lines:
        data_cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if len(data_cells) > 0:
            row = table.add_row()
            for i, cell_text in enumerate(data_cells):
                if i < len(row.cells):
                    cell = row.cells[i]
                    # Pulisci il paragrafo esistente
                    cell.paragraphs[0].clear()
                    # Aggiungi testo formattato (gestisce markdown nelle celle)
                    add_formatted_text(cell.paragraphs[0], cell_text)
    
    # Aggiungi spazio dopo la tabella
    doc.add_paragraph()

def read_file_content(file_path):
    """Legge il contenuto di un file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Errore nella lettura del file {file_path}: {e}")
        return ""

def main():
    """Funzione principale"""
    
    # Verifica che la cartella Relazione_4808 esista
    base_dir = Path("Relazione_4808")
    if not base_dir.exists():
        print("Errore: La cartella 'Relazione_4808' non esiste!")
        print("Assicurati di eseguire lo script dalla directory corretta.")
        sys.exit(1)
    
    # Crea il documento Word
    doc = Document()
    
    # Configura gli stili
    setup_document_styles(doc)
    
    # Titolo principale del documento
    doc.add_paragraph("RELAZIONE VIAGGIO 4808", style='CustomTitle')
    doc.add_paragraph("CANADA GRANDE NORD 2025", style='CustomTitle')
    doc.add_paragraph("03 - 24 Agosto 2025", style='CustomTitle')
    
    add_page_break(doc)
    
    # Lista dei file delle sezioni principali in ordine
    main_sections = [
        "00_Indice.md",
        "01_Informazioni_Generali.md",
        "02_Considerazioni_Iniziali.md",
        "03_Informazioni_Pratiche.md",
        "04_Piano_Voli.md",
        "05_Trasporti.md",
        "06_Pernottamenti.md",
        "07_Costi_e_Cassa.md",
        "08_Escursioni_e_Attivita.md"
    ]
    
    # Aggiungi le sezioni principali
    print("Aggiungendo sezioni principali...")
    for section_file in main_sections:
        file_path = base_dir / section_file
        if file_path.exists():
            print(f"  - {section_file}")
            content = read_file_content(file_path)
            if content:
                process_markdown_content(doc, content)
                add_page_break(doc)
        else:
            print(f"  - ATTENZIONE: {section_file} non trovato!")
    
    # Aggiungi titolo per l'itinerario giornaliero
    doc.add_paragraph("ITINERARIO GIORNALIERO DETTAGLIATO", style='SectionTitle')
    add_page_break(doc)
    
    # Aggiungi i giorni dell'itinerario
    print("Aggiungendo itinerario giornaliero...")
    itinerario_dir = base_dir / "Itinerario_Giornaliero"
    
    if itinerario_dir.exists():
        # Ordina i file per giorno
        day_files = sorted([f for f in itinerario_dir.glob("Giorno_*.md")])
        
        for day_file in day_files:
            print(f"  - {day_file.name}")
            content = read_file_content(day_file)
            if content:
                process_markdown_content(doc, content)
                add_page_break(doc)
    else:
        print("  - ATTENZIONE: Cartella Itinerario_Giornaliero non trovata!")
    
    # Aggiungi il manuale operativo
    print("Aggiungendo manuale operativo coordinatore...")
    manual_file = base_dir / "MANUALE_OPERATIVO_COORDINATORE.md"
    if manual_file.exists():
        content = read_file_content(manual_file)
        if content:
            process_markdown_content(doc, content)
    else:
        print("  - ATTENZIONE: MANUALE_OPERATIVO_COORDINATORE.md non trovato!")
    
    # Salva il documento
    output_file = "Relazione_4808_Canada_Grande_Nord_2025.docx"
    doc.save(output_file)
    
    print(f"\n✅ Documento Word creato con successo: {output_file}")
    print(f"📄 Il file contiene:")
    print(f"   - Tutte le sezioni principali della relazione")
    print(f"   - Itinerario dettagliato di tutti i 22 giorni")
    print(f"   - Manuale operativo per il coordinatore")
    print(f"\n🎯 Il documento è pronto per essere utilizzato!")

if __name__ == "__main__":
    # Verifica che python-docx sia installato
    try:
        from docx import Document
    except ImportError:
        print("❌ Errore: Il modulo 'python-docx' non è installato!")
        print("\n📦 Per installarlo, esegui:")
        print("   pip install python-docx")
        print("\n🔄 Poi rilancia questo script.")
        sys.exit(1)
    
    main()
